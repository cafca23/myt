import streamlit as st
import google.generativeai as genai
import pdfplumber
import requests
import datetime
import urllib.parse
import pandas as pd
import time
import urllib3
from docx import Document
from io import BytesIO

urllib3.disable_warnings(urllib3.exceptions.InsecureRequestWarning)

# ==========================================
# ⚙️ 1. 기본 설정 및 API 키
# ==========================================
st.set_page_config(page_title="마이티시스템 입찰 플랫폼 V3", layout="wide")

GEMINI_API_KEY = st.secrets["GEMINI_API_KEY"]
KONEPS_API_KEY = "fc9942134c063694eeb5dad340a314eec93995f86031e3653cddb5d4d38dfbd3"

genai.configure(api_key=GEMINI_API_KEY)
model = genai.GenerativeModel('gemini-2.5-flash')

# ==========================================
# 👤 2. 마이티시스템 V3 프로필 (★출력 양식 완전 개편★)
# ==========================================
mighty_profile = """
당신은 '마이티시스템'의 입찰 전문 전략 비서입니다. 
제공된 공고문과 과업지시서를 분석하여 아래의 지정된 [출력 양식]에 맞춰 한 치의 오차 없이 보고서를 작성하십시오.

[우리 회사 자격 요건]
1. 기본: 소기업/소상공인, 신용등급 BBO
2. 면허: 정보통신공사업(0036), SW사업자(1426, 1468, 1470)
3. 시평액: 5,595,600,000원
4. 직생 보유: 8111229901, 8111159801, 8111159901, 8111179901, 8111181101, 8111189901 등
5. 실적: ★현재 수행 실적 전혀 없음(0건)★. 실적 제한 공고는 무조건 '불가능' 처리할 것.

[출력 양식] - 반드시 아래의 5가지 목차와 양식을 지킬 것.

### 1. 📋 공고 기본 정보
- **공고명:** - **입찰공고번호:** - **수요기관:** - **배정예산:** ### 2. 🚨 입찰 참가자격 정밀 진단 (★가장 중요)
- 공고문에 명시된 참가자격을 모두 찾아내어 아래 마크다운 표 형식으로 작성할 것.
| 공고문 요구 자격 | 당사 충족 여부 | 판정 사유 및 상세 설명 (불가능 시 이유 명시) |
| :--- | :---: | :--- |
| (예: 정보통신공사업 등록업체) | 🟢 가능 | 당사 정보통신공사업(0036) 면허 보유 |
| (예: 최근 3년 내 1억 이상 실적) | 🔴 불가능 | 당사 현재 실적 0건으로 자격 미달 |

### 3. 📑 제안서 작성 주요 내용
- 제안서 제출이 필수인지 확인하고, 필수라면 과업지시서/제안요청서(RFP)를 바탕으로 우리가 제안서에 반드시 포함해야 할 핵심 목차나 기술적 요구사항을 요약할 것. (※ 제안서 평가가 없는 단순 최저가 입찰의 경우 '제안서 작성 불필요'로 명시)

### 4. ⚖️ 평가 및 협상적격자 선정 방법
- **평가 비율:** 기술평가 00% / 가격평가 00% (비율 명시)
- **선정 기준:** 기술능력평가점수 배점한도의 00% 이상인 자 등 커트라인 명시
- **★실적 평가 분석:** 정량평가(수행실적) 배점이 몇 점인지 찾아내어 명시할 것. (당사는 실적이 0건이므로, 실적에서 몇 점을 감점받고 시작하는지 객관적으로 분석할 것)

### 5. ⚠️ 위험 요소 및 특이사항
- 과업지시서 상의 독소 조항, 무상 유지보수 기간, 납기일, 지체상금 등 주의사항 요약.
"""

# 보고서 생성 함수
def create_report(analysis_text):
    doc = Document()
    doc.add_heading('마이티시스템 입찰 적격성 및 제안 분석 보고서', 0)
    doc.add_paragraph(f"분석 일시: {datetime.datetime.now().strftime('%Y-%m-%d %H:%M')}")
    doc.add_paragraph("-" * 50)
    doc.add_paragraph(analysis_text)
    
    bio = BytesIO()
    doc.save(bio)
    return bio.getvalue()

# ==========================================
# 🌐 3. 메인 UI (검색 및 분석)
# ==========================================
st.title("🚀 마이티시스템 입찰 플랫폼 V3 (표준 보고서형)")
st.markdown("---")

# --- 1단계: 검색 ---
st.header("📊 1단계: 나라장터 실시간 맞춤 검색")
if "bids_df" not in st.session_state: st.session_state.bids_df = None

if st.button("🔍 맞춤 공고 검색 시작", type="primary"):
    with st.status("공고 수집 중...", expanded=False) as status:
        keywords = ['정보시스템', '전산시스템', '서버', '스토리지', '인프라', '유지관리']
        all_bids = []
        start_date = (datetime.datetime.now() - datetime.timedelta(days=7)).strftime('%Y%m%d0000')
        end_date = datetime.datetime.now().strftime('%Y%m%d2359')
        urls = {"물품": "https://apis.data.go.kr/1230000/ad/BidPublicInfoService/getBidPblancListInfoThngPPSSrch",
                "용역": "https://apis.data.go.kr/1230000/ad/BidPublicInfoService/getBidPblancListInfoServcPPSSrch"}
        
        for kind, base_url in urls.items():
            for kw in keywords:
                encoded_kw = urllib.parse.quote(kw)
                req_url = f"{base_url}?ServiceKey={KONEPS_API_KEY}&numOfRows=30&pageNo=1&inqryDiv=1&inqryBgnDt={start_date}&inqryEndDt={end_date}&bidNtceNm={encoded_kw}&type=json"
                try:
                    res = requests.get(req_url, headers={'User-Agent': 'Mozilla/5.0'}, timeout=15, verify=False)
                    if res.status_code == 200:
                        items = res.json()['response']['body'].get('items', [])
                        for item in items:
                            all_bids.append({
                                '상세링크': item.get('bidNtceDtlUrl', ''), '분류': kind, '공고명': item.get('bidNtceNm', ''),
                                '수요기관': item.get('ntceInsttNm', ''), '마감일시': item.get('bidClseDt', ''), '공고번호': item.get('bidNtceNo', '')
                            })
                except: pass
        
        if all_bids:
            df = pd.DataFrame(all_bids).drop_duplicates(subset=['공고번호']).sort_values(by='마감일시')
            st.session_state.bids_df = df
            status.update(label="✅ 검색 완료", state="complete")

if st.session_state.bids_df is not None:
    st.dataframe(st.session_state.bids_df, use_container_width=True, column_config={"상세링크": st.column_config.LinkColumn("상세보기", display_text="👉 이동")})

st.markdown("<br><hr><br>", unsafe_allow_html=True)

# --- 2단계: 심층 분석 ---
st.header("🤖 2단계: AI 입찰/제안 심층 분석")
st.info("💡 입찰참가자격을 표로 정리하고, 제안서 요건과 평가 방식을 분석합니다.")

col1, col2 = st.columns(2)
with col1:
    notice_file = st.file_uploader("1️⃣ 입찰공고서 업로드 (필수)", type="pdf")
with col2:
    rfp_file = st.file_uploader("2️⃣ 제안요청서/과업지시서 업로드 (선택)", type="pdf")

if notice_file:
    if st.button("🧐 마이티시스템 맞춤형 리포트 생성"):
        with st.spinner("참가자격을 표로 대조하고 평가 기준을 분석 중입니다..."):
            full_text = ""
            for f in [notice_file, rfp_file]:
                if f:
                    with pdfplumber.open(f) as pdf:
                        for page in pdf.pages:
                            full_text += page.extract_text() + "\n"
            
            # AI 분석 실행
            response = model.generate_content(mighty_profile + "\n\n[입찰문서 원본 내용]\n" + full_text)
            st.session_state.analysis_result = response.text
            
            st.success("✅ 분석 완료!")
            with st.container(border=True):
                st.markdown(st.session_state.analysis_result)
            
            # 워드 리포트 생성
            report_data = create_report(st.session_state.analysis_result)
            st.download_button(
                label="📥 분석 결과 보고서(Word) 다운로드",
                data=report_data,
                file_name=f"마이티시스템_입찰분석_{datetime.date.today()}.docx",
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
            )
